"""
Template-aware DOCX generators for all 15 TailorMyCV front-end templates.

Public API
----------
    generate_docx_from_key(resume_data, template_key, bold_keywords) -> bytes

Each of the 15 template keys is mapped to a TemplateConfig that controls:
  layout  — single | sidebar | two-equal
  header  — centered | banner | serif-centered | left
  heading — rule | colored | left-border | double-rule | gold-rule
  accent  — primary hex colour (headings, rules, sidebar bg)
  font    — Calibri | Times New Roman | Georgia
  compact — tighter spacing for 1-page templates

Three base renderers cover every combination:
  _render_single    — single-column (12 templates)
  _render_sidebar   — coloured left-sidebar (Prism, Vivid)
  _render_symmetry  — full-width header + two equal columns (Symmetry)
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


# ═══════════════════════════════════════════════════════════════════════════════
# Low-level XML / style helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_width_twips(cell, twips: int) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_cell_padding(cell, top=100, left=170, bottom=100, right=170) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _remove_table_borders(table) -> None:
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_table_width_twips(table, twips: int) -> None:
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _para_border(para, color: str, pos: str = "bottom", size: int = 6) -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement(f"w:{pos}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), "4")
    b.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(b)
    pPr.append(pBdr)


def _para_left_border(para, color: str, size: int = 18) -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:left")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), "6")
    b.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(b)
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════════════════
# Template config
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class TemplateConfig:
    accent: str           # hex, no '#'
    header: str           # centered | banner | serif-centered | left
    font: str             # font family name
    heading: str          # rule | colored | left-border | double-rule | gold-rule | circle-marker
    compact: bool         # reduced spacing for 1-page dense layouts
    layout: str           # single | sidebar | two-equal | left-bar
    sidebar_color: str    # sidebar background hex (sidebar / two-equal / left-bar layouts)
    sidebar_ratio: float  # sidebar share of content width (e.g. 0.32)
    banner_bg: str = ""   # override banner background colour (defaults to accent)


_MM_TO_TWIPS = 56.7
_CONTENT_W_MM = 180.0   # A4 210mm − 15mm margins each side


_CONFIGS: dict[str, TemplateConfig] = {
    "Cambridge": TemplateConfig(
        accent="1f2937", header="centered",      font="Calibri",          heading="rule",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "Horizon": TemplateConfig(
        accent="1d4ed8", header="banner",        font="Calibri",          heading="rule",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "Prestige": TemplateConfig(
        accent="374151", header="serif-centered",font="Times New Roman",   heading="double-rule",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "Admiral": TemplateConfig(
        accent="1e3a5f", header="banner",        font="Calibri",          heading="rule",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "Swift": TemplateConfig(
        accent="1e293b", header="left",          font="Calibri",          heading="colored",
        compact=True,  layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "Catalyst": TemplateConfig(
        accent="ea580c", header="left",          font="Calibri",          heading="left-border",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "Canvas": TemplateConfig(
        accent="9ca3af", header="centered",      font="Calibri",          heading="rule",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "Jade": TemplateConfig(
        accent="0d9488", header="banner",        font="Calibri",          heading="rule",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "Prism": TemplateConfig(
        accent="2563eb", header="centered",      font="Calibri",          heading="rule",
        compact=False, layout="sidebar",  sidebar_color="2563eb", sidebar_ratio=0.30,
    ),
    "Vivid": TemplateConfig(
        accent="7c3aed", header="centered",      font="Calibri",          heading="rule",
        compact=False, layout="sidebar",  sidebar_color="7c3aed", sidebar_ratio=0.30,
    ),
    "Chronicle": TemplateConfig(
        accent="1d4ed8", header="left",          font="Calibri",          heading="rule",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "Summit": TemplateConfig(
        accent="0f172a", header="banner",        font="Calibri",          heading="colored",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "Symmetry": TemplateConfig(
        accent="1e3a5f", header="centered",      font="Calibri",          heading="rule",
        compact=False, layout="two-equal",sidebar_color="f1f5f9", sidebar_ratio=0.50,
    ),
    "Scholar": TemplateConfig(
        accent="374151", header="left",          font="Times New Roman",   heading="rule",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "Luxe": TemplateConfig(
        accent="b45309", header="serif-centered",font="Georgia",           heading="gold-rule",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    # ── New templates ──────────────────────────────────────────────────────────
    "TechModern": TemplateConfig(
        accent="10b981", header="banner",        font="Courier New",       heading="colored",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
        banner_bg="0f172a",   # dark banner, green accent headings
    ),
    "Pulse": TemplateConfig(
        accent="e11d48", header="left",          font="Calibri",           heading="rule",
        compact=False, layout="left-bar", sidebar_color="e11d48", sidebar_ratio=0.035,
    ),
    "HexagonPro": TemplateConfig(
        accent="0ea5e9", header="centered",      font="Calibri",           heading="circle-marker",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "SalesImpact": TemplateConfig(
        accent="dc2626", header="banner",        font="Calibri",           heading="rule",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
    "Healthcare": TemplateConfig(
        accent="0891b2", header="centered",      font="Calibri",           heading="left-border",
        compact=False, layout="single",   sidebar_color="", sidebar_ratio=0.0,
    ),
}


# ═══════════════════════════════════════════════════════════════════════════════
# Paragraph / run helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _sp(para, before: float = 0, after: float = 4) -> None:
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)


def _add_run(para, text: str, *, size: float = 10, bold: bool = False,
             italic: bool = False, color: str | None = None, font: str = "Calibri"):
    r = para.add_run(text)
    r.font.size  = Pt(size)
    r.bold       = bold
    r.italic     = italic
    r.font.name  = font
    if color:
        r.font.color.rgb = _rgb(color)
    return r


def _contact_str(contact: dict) -> str:
    parts = []
    for key in ("email", "phone", "location", "linkedin"):
        val = contact.get(key, "")
        if not val:
            continue
        if key == "linkedin":
            val = re.sub(r"^https?://(www\.)?", "", val).rstrip("/")
        parts.append(val)
    return "  |  ".join(parts)


# ═══════════════════════════════════════════════════════════════════════════════
# Single-column renderer
# ═══════════════════════════════════════════════════════════════════════════════

def _render_single(r: dict, cfg: TemplateConfig) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Mm(210)
    sec.page_height   = Mm(297)
    m = 10 if cfg.compact else 15
    sec.left_margin   = Mm(m)
    sec.right_margin  = Mm(m)
    sec.top_margin    = Mm(10 if cfg.compact else 14)
    sec.bottom_margin = Mm(10 if cfg.compact else 12)

    contact = r.get("contact", {})
    fn      = cfg.font
    ac      = cfg.accent
    body_pt = 9.5 if cfg.compact else 10
    gap_pt  = 1.5 if cfg.compact else 4

    # ── Header ────────────────────────────────────────────────────────────────
    if cfg.header == "banner":
        banner_color = cfg.banner_bg if cfg.banner_bg else ac
        tbl = doc.add_table(rows=1, cols=1)
        _remove_table_borders(tbl)
        _set_table_width_twips(tbl, int(_CONTENT_W_MM * _MM_TO_TWIPS))
        cell = tbl.cell(0, 0)
        _set_cell_bg(cell, banner_color)
        _set_cell_padding(cell, top=200, left=300, bottom=200, right=300)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p, after=4)
        _add_run(p, (r.get("name") or "").upper(), size=22, bold=True, color="ffffff", font=fn)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p2, after=0)
        _add_run(p2, _contact_str(contact), size=9, color="d1d5db", font=fn)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    elif cfg.header == "centered":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p, after=2)
        _add_run(p, r.get("name", ""), size=22, bold=True, color=ac, font=fn)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p2, after=8)
        _add_run(p2, _contact_str(contact), size=9, color="6b7280", font=fn)

    elif cfg.header == "serif-centered":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p, before=4, after=4)
        _para_border(p, ac, "top", 10)
        _add_run(p, r.get("name", ""), size=24, bold=True, color=ac, font=fn)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p2, after=6)
        _para_border(p2, ac, "bottom", 6)
        _add_run(p2, _contact_str(contact), size=9, color="6b7280", font=fn)

    else:  # left
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _sp(p, after=2)
        _add_run(p, r.get("name", ""), size=20, bold=True, color=ac, font=fn)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _sp(p2, after=8)
        _add_run(p2, _contact_str(contact), size=9, color="6b7280", font=fn)

    # ── Section heading factory ───────────────────────────────────────────────
    def _heading(title: str):
        p = doc.add_paragraph()
        before = 6 if cfg.compact else 10
        _sp(p, before=before, after=2)
        if cfg.heading == "rule":
            _para_border(p, ac, "bottom", 6)
            _add_run(p, title.upper(), size=10, bold=True, color=ac, font=fn)
        elif cfg.heading == "colored":
            _add_run(p, title.upper(), size=10, bold=True, color=ac, font=fn)
        elif cfg.heading == "left-border":
            _para_left_border(p, ac, size=18)
            p.paragraph_format.left_indent = Pt(6)
            _add_run(p, title.upper(), size=10, bold=True, color=ac, font=fn)
        elif cfg.heading == "double-rule":
            _para_border(p, ac, "top", 8)
            _para_border(p, ac, "bottom", 4)
            _add_run(p, title.upper(), size=10, bold=True, color=ac, font=fn)
        elif cfg.heading == "gold-rule":
            _para_border(p, ac, "bottom", 8)
            _add_run(p, title.upper(), size=10, bold=True, color=ac, font=fn)
        elif cfg.heading == "circle-marker":
            _add_run(p, "◉  ", size=10, bold=True, color=ac, font=fn)
            _add_run(p, title.upper(), size=10, bold=True, color="0f172a", font=fn)
            _para_border(p, "e2e8f0", "bottom", 4)

    # ── Bullet ────────────────────────────────────────────────────────────────
    def _bullet(text: str):
        p = doc.add_paragraph()
        _sp(p, after=1)
        run = p.add_run(f"•  {text}")
        run.font.size = Pt(body_pt)
        run.font.name = fn

    # ── Sections ──────────────────────────────────────────────────────────────
    if r.get("summary"):
        _heading("Professional Summary")
        p = doc.add_paragraph(r["summary"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _sp(p, after=gap_pt)
        for run in p.runs:
            run.font.size = Pt(body_pt)
            run.font.name = fn

    if r.get("experience"):
        _heading("Experience")
        for job in r["experience"]:
            p = doc.add_paragraph()
            _sp(p, after=1)
            _add_run(p, job.get("role", ""), size=10.5, bold=True, color="111827", font=fn)
            _add_run(p, f"  ·  {job.get('company', '')}", size=10.5, font=fn)
            p2 = doc.add_paragraph()
            _sp(p2, after=2)
            _add_run(p2, job.get("dates", ""), size=9, italic=True, color="6b7280", font=fn)
            for b in job.get("bullets", []):
                _bullet(b)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    if r.get("education"):
        _heading("Education")
        for ed in r["education"]:
            p = doc.add_paragraph()
            _sp(p, after=1)
            _add_run(p, ed.get("degree", ""), size=10.5, bold=True, color="111827", font=fn)
            _add_run(p, f"  ·  {ed.get('institution', '')}", size=10.5, font=fn)
            p2 = doc.add_paragraph()
            _sp(p2, after=gap_pt)
            _add_run(p2, ed.get("dates", ""), size=9, italic=True, color="6b7280", font=fn)

    for sec_data in r.get("sections", []):
        title = sec_data.get("title", "")
        items = sec_data.get("items", [])
        if title and items:
            _heading(title)
            for item in items:
                _bullet(item)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# Sidebar renderer  (Prism, Vivid)
# Left column: coloured background — name, contact, sections/skills
# Right column: white — summary, experience, education
# ═══════════════════════════════════════════════════════════════════════════════

def _render_sidebar(r: dict, cfg: TemplateConfig) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Mm(210)
    sec.page_height   = Mm(297)
    sec.left_margin   = Mm(0)
    sec.right_margin  = Mm(0)
    sec.top_margin    = Mm(0)
    sec.bottom_margin = Mm(10)

    contact  = r.get("contact", {})
    fn       = cfg.font
    ac       = cfg.accent
    sc       = cfg.sidebar_color

    PAGE_W_TWIPS  = int(210 * _MM_TO_TWIPS)
    SIDE_TWIPS    = int(PAGE_W_TWIPS * cfg.sidebar_ratio)
    MAIN_TWIPS    = PAGE_W_TWIPS - SIDE_TWIPS

    tbl = doc.add_table(rows=1, cols=2)
    _remove_table_borders(tbl)
    _set_table_width_twips(tbl, PAGE_W_TWIPS)
    tbl.allow_autofit = False

    left_cell  = tbl.cell(0, 0)
    right_cell = tbl.cell(0, 1)
    _set_cell_width_twips(left_cell,  SIDE_TWIPS)
    _set_cell_width_twips(right_cell, MAIN_TWIPS)
    _set_cell_bg(left_cell, sc)
    _set_cell_padding(left_cell,  top=300, left=250, bottom=300, right=200)
    _set_cell_padding(right_cell, top=240, left=280, bottom=200, right=280)

    # ── Sidebar: name ─────────────────────────────────────────────────────────
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _sp(lp, before=0, after=6)
    _add_run(lp, (r.get("name") or "").upper(), size=13, bold=True, color="ffffff", font=fn)

    def _side_section(title: str):
        p = left_cell.add_paragraph()
        _sp(p, before=12, after=3)
        _para_border(p, "ffffff", "bottom", 4)
        _add_run(p, title.upper(), size=8, bold=True, color="ffffff", font=fn)

    def _side_item(text: str):
        p = left_cell.add_paragraph()
        _sp(p, after=2)
        _add_run(p, text, size=8.5, color="dbeafe", font=fn)

    # Sidebar: contact
    _side_section("Contact")
    for key in ("email", "phone", "location", "linkedin"):
        val = contact.get(key, "")
        if val:
            if key == "linkedin":
                val = re.sub(r"^https?://(www\.)?", "", val).rstrip("/")
            _side_item(val)

    # Sidebar: additional sections (skills, etc.)
    for sec_data in r.get("sections", []):
        title = sec_data.get("title", "")
        items = sec_data.get("items", [])
        if title and items:
            _side_section(title)
            for item in items:
                _side_item(item)

    # ── Main column ───────────────────────────────────────────────────────────
    def _main_heading(title: str, container):
        p = container.add_paragraph()
        _sp(p, before=10, after=2)
        _para_border(p, ac, "bottom", 6)
        _add_run(p, title.upper(), size=10, bold=True, color=ac, font=fn)

    def _main_bullet(text: str, container):
        p = container.add_paragraph()
        _sp(p, after=1)
        run = p.add_run(f"•  {text}")
        run.font.size = Pt(10)
        run.font.name = fn

    # Clear the auto-created first paragraph
    rp = right_cell.paragraphs[0]
    _sp(rp, after=0)

    if r.get("summary"):
        _main_heading("Professional Summary", right_cell)
        p = right_cell.add_paragraph(r["summary"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _sp(p, after=4)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = fn

    if r.get("experience"):
        _main_heading("Experience", right_cell)
        for job in r["experience"]:
            p = right_cell.add_paragraph()
            _sp(p, after=1)
            _add_run(p, job.get("role", ""), size=10.5, bold=True, color="111827", font=fn)
            _add_run(p, f"  ·  {job.get('company', '')}", size=10, font=fn)
            p2 = right_cell.add_paragraph()
            _sp(p2, after=2)
            _add_run(p2, job.get("dates", ""), size=9, italic=True, color="6b7280", font=fn)
            for b in job.get("bullets", []):
                _main_bullet(b, right_cell)
            right_cell.add_paragraph().paragraph_format.space_after = Pt(2)

    if r.get("education"):
        _main_heading("Education", right_cell)
        for ed in r["education"]:
            p = right_cell.add_paragraph()
            _sp(p, after=1)
            _add_run(p, ed.get("degree", ""), size=10.5, bold=True, color="111827", font=fn)
            _add_run(p, f"  ·  {ed.get('institution', '')}", size=10, font=fn)
            p2 = right_cell.add_paragraph()
            _sp(p2, after=4)
            _add_run(p2, ed.get("dates", ""), size=9, italic=True, color="6b7280", font=fn)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# Two-equal-column renderer  (Symmetry)
# Full-width header, then left=Experience / right=Summary+Skills+Education
# ═══════════════════════════════════════════════════════════════════════════════

def _render_symmetry(r: dict, cfg: TemplateConfig) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Mm(210)
    sec.page_height   = Mm(297)
    sec.left_margin   = Mm(15)
    sec.right_margin  = Mm(15)
    sec.top_margin    = Mm(14)
    sec.bottom_margin = Mm(12)

    contact = r.get("contact", {})
    fn      = cfg.font
    ac      = cfg.accent
    sc      = cfg.sidebar_color

    # ── Full-width header ─────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, after=2)
    _para_border(p, ac, "bottom", 8)
    _add_run(p, r.get("name", ""), size=22, bold=True, color=ac, font=fn)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p2, after=8)
    _add_run(p2, _contact_str(contact), size=9, color="6b7280", font=fn)

    # ── Two-column body table ─────────────────────────────────────────────────
    CONTENT_W_TWIPS = int(_CONTENT_W_MM * _MM_TO_TWIPS)
    HALF_TWIPS      = CONTENT_W_TWIPS // 2
    GAP_TWIPS       = int(4 * _MM_TO_TWIPS)

    tbl = doc.add_table(rows=1, cols=2)
    _remove_table_borders(tbl)
    _set_table_width_twips(tbl, CONTENT_W_TWIPS)
    tbl.allow_autofit = False

    left_cell  = tbl.cell(0, 0)
    right_cell = tbl.cell(0, 1)
    _set_cell_width_twips(left_cell,  HALF_TWIPS - GAP_TWIPS // 2)
    _set_cell_width_twips(right_cell, HALF_TWIPS - GAP_TWIPS // 2)
    _set_cell_padding(left_cell,  top=0, left=0, bottom=0, right=200)
    _set_cell_padding(right_cell, top=0, left=200, bottom=0, right=0)
    _set_cell_bg(right_cell, sc)

    def _col_heading(title: str, container):
        p = container.add_paragraph()
        _sp(p, before=8, after=2)
        _para_border(p, ac, "bottom", 6)
        _add_run(p, title.upper(), size=10, bold=True, color=ac, font=fn)

    def _col_bullet(text: str, container):
        p = container.add_paragraph()
        _sp(p, after=1)
        run = p.add_run(f"•  {text}")
        run.font.size = Pt(10)
        run.font.name = fn

    # Left column: Experience
    lp0 = left_cell.paragraphs[0]
    _sp(lp0, after=0)

    if r.get("experience"):
        _col_heading("Experience", left_cell)
        for job in r["experience"]:
            p = left_cell.add_paragraph()
            _sp(p, after=1)
            _add_run(p, job.get("role", ""), size=10.5, bold=True, color="111827", font=fn)
            _add_run(p, f"  ·  {job.get('company', '')}", size=10, font=fn)
            p2 = left_cell.add_paragraph()
            _sp(p2, after=2)
            _add_run(p2, job.get("dates", ""), size=9, italic=True, color="6b7280", font=fn)
            for b in job.get("bullets", []):
                _col_bullet(b, left_cell)
            left_cell.add_paragraph().paragraph_format.space_after = Pt(2)

    # Right column: Summary + sections + Education
    rp0 = right_cell.paragraphs[0]
    _sp(rp0, after=0)
    _set_cell_padding(right_cell, top=0, left=240, bottom=0, right=0)

    if r.get("summary"):
        _col_heading("Summary", right_cell)
        p = right_cell.add_paragraph(r["summary"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _sp(p, after=4)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = fn

    for sec_data in r.get("sections", []):
        title = sec_data.get("title", "")
        items = sec_data.get("items", [])
        if title and items:
            _col_heading(title, right_cell)
            for item in items:
                _col_bullet(item, right_cell)

    if r.get("education"):
        _col_heading("Education", right_cell)
        for ed in r["education"]:
            p = right_cell.add_paragraph()
            _sp(p, after=1)
            _add_run(p, ed.get("degree", ""), size=10.5, bold=True, color="111827", font=fn)
            _add_run(p, f"  ·  {ed.get('institution', '')}", size=10, font=fn)
            p2 = right_cell.add_paragraph()
            _sp(p2, after=4)
            _add_run(p2, ed.get("dates", ""), size=9, italic=True, color="6b7280", font=fn)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# Left-bar renderer  (Pulse)
# Narrow decorative left bar (no content) + all content in right column
# ═══════════════════════════════════════════════════════════════════════════════

def _render_left_bar(r: dict, cfg: TemplateConfig) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Mm(210)
    sec.page_height   = Mm(297)
    sec.left_margin   = Mm(0)
    sec.right_margin  = Mm(0)
    sec.top_margin    = Mm(0)
    sec.bottom_margin = Mm(10)

    contact  = r.get("contact", {})
    fn       = cfg.font
    ac       = cfg.accent
    sc       = cfg.sidebar_color

    BAR_MM     = int(210 * cfg.sidebar_ratio)    # narrow coloured bar
    CONTENT_MM = 210 - BAR_MM
    PAGE_TWIPS = int(210 * _MM_TO_TWIPS)
    BAR_TWIPS  = int(BAR_MM * _MM_TO_TWIPS)
    CON_TWIPS  = int(CONTENT_MM * _MM_TO_TWIPS)

    tbl = doc.add_table(rows=1, cols=2)
    _remove_table_borders(tbl)
    _set_table_width_twips(tbl, PAGE_TWIPS)
    tbl.allow_autofit = False

    bar_cell  = tbl.cell(0, 0)
    main_cell = tbl.cell(0, 1)
    _set_cell_width_twips(bar_cell,  BAR_TWIPS)
    _set_cell_width_twips(main_cell, CON_TWIPS)
    _set_cell_bg(bar_cell, sc)
    _set_cell_padding(bar_cell,  top=0, left=0, bottom=0, right=0)
    _set_cell_padding(main_cell, top=240, left=300, bottom=200, right=300)

    # Bar: empty — just the background colour stretches to content height
    bp = bar_cell.paragraphs[0]
    _sp(bp, after=0)

    # Main column: full resume content
    def _heading(title: str):
        p = main_cell.add_paragraph()
        _sp(p, before=8, after=2)
        _para_border(p, ac, "bottom", 6)
        _add_run(p, title.upper(), size=10, bold=True, color=ac, font=fn)

    def _bullet(text: str):
        p = main_cell.add_paragraph()
        _sp(p, after=1)
        run = p.add_run(f"•  {text}")
        run.font.size = Pt(10)
        run.font.name = fn

    # Name header in main column
    p0 = main_cell.paragraphs[0]
    _sp(p0, before=0, after=2)
    _add_run(p0, r.get("name", ""), size=22, bold=True, color=ac, font=fn)
    p1 = main_cell.add_paragraph()
    _sp(p1, after=2)
    _add_run(p1, r.get("contact", {}).get("email", "") or "", size=9, color="6b7280", font=fn)
    p2 = main_cell.add_paragraph()
    _sp(p2, after=8)
    _add_run(p2, _contact_str(contact), size=9, color="6b7280", font=fn)

    if r.get("summary"):
        _heading("Professional Summary")
        p = main_cell.add_paragraph(r["summary"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _sp(p, after=4)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = fn

    if r.get("experience"):
        _heading("Experience")
        for job in r["experience"]:
            p = main_cell.add_paragraph()
            _sp(p, after=1)
            _add_run(p, job.get("role", ""), size=10.5, bold=True, color="111827", font=fn)
            _add_run(p, f"  ·  {job.get('company', '')}", size=10.5, font=fn)
            p2 = main_cell.add_paragraph()
            _sp(p2, after=2)
            _add_run(p2, job.get("dates", ""), size=9, italic=True, color="6b7280", font=fn)
            for b in job.get("bullets", []):
                _bullet(b)
            main_cell.add_paragraph().paragraph_format.space_after = Pt(2)

    if r.get("education"):
        _heading("Education")
        for ed in r["education"]:
            p = main_cell.add_paragraph()
            _sp(p, after=1)
            _add_run(p, ed.get("degree", ""), size=10.5, bold=True, color="111827", font=fn)
            _add_run(p, f"  ·  {ed.get('institution', '')}", size=10, font=fn)
            p2 = main_cell.add_paragraph()
            _sp(p2, after=4)
            _add_run(p2, ed.get("dates", ""), size=9, italic=True, color="6b7280", font=fn)

    for sec_data in r.get("sections", []):
        title = sec_data.get("title", "")
        items = sec_data.get("items", [])
        if title and items:
            _heading(title)
            for item in items:
                _bullet(item)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# Public entry point
# ═══════════════════════════════════════════════════════════════════════════════

KNOWN_TEMPLATE_KEYS: frozenset[str] = frozenset(_CONFIGS.keys())


def generate_docx_from_key(
    resume_data: dict,
    template_key: str,
    bold_keywords: list[str] | None = None,   # reserved — future keyword bolding
) -> bytes:
    """Generate a styled DOCX for the given template key.

    Falls back to the single-column Cambridge style when the key is unknown.
    """
    cfg = _CONFIGS.get(template_key, _CONFIGS["Cambridge"])

    if cfg.layout == "sidebar":
        return _render_sidebar(resume_data, cfg)
    if cfg.layout == "two-equal":
        return _render_symmetry(resume_data, cfg)
    if cfg.layout == "left-bar":
        return _render_left_bar(resume_data, cfg)
    return _render_single(resume_data, cfg)
