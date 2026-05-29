"""Generate the three prebuilt DOCX resume templates.

Run from the backend directory:
    python scripts/generate_templates.py

Overwrites templates/prebuilt/clean.docx, modern.docx, executive.docx
"""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates", "prebuilt")

# ── Shared colours ─────────────────────────────────────────────────────────────

BLACK     = RGBColor(0x1A, 0x1A, 0x1A)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY  = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY= RGBColor(0x88, 0x88, 0x88)
BRAND     = RGBColor(0x2B, 0x57, 0x9A)   # #2B579A
TEAL      = RGBColor(0x0D, 0x94, 0x88)   # #0D9488
ACCENT_LT = RGBColor(0xE8, 0xF0, 0xFB)   # light brand tint


# ── Helpers ────────────────────────────────────────────────────────────────────

def _run(para, text: str, font: str, size: float, bold=False, italic=False,
         color: RGBColor = None, underline=False):
    r = para.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    if color:
        r.font.color.rgb = color
    return r


def _para(doc, text: str = "", font: str = "Calibri", size: float = 10,
          bold=False, italic=False, color: RGBColor = None,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before: float = 0, space_after: float = 4) -> object:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        _run(p, text, font, size, bold=bold, italic=italic, color=color)
    return p


def _hr(doc, color: RGBColor = None, thickness_pt: float = 0.5,
        space_before: float = 0, space_after: float = 3):
    """Add a horizontal rule using paragraph border XML."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(thickness_pt * 8)))
    bottom.set(qn("w:space"), "1")
    hex_color = "auto" if color is None else f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _double_hr(doc, color: RGBColor = None, space_before: float = 0, space_after: float = 3):
    """Add a double horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom"):
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:val"), "double")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "1")
        hex_color = "auto" if color is None else f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
        elem.set(qn("w:color"), hex_color)
        pBdr.append(elem)
    pPr.append(pBdr)
    return p


def _set_margins(doc, top=0.8, bottom=0.8, left=0.9, right=0.9):
    for sec in doc.sections:
        sec.top_margin    = Inches(top)
        sec.bottom_margin = Inches(bottom)
        sec.left_margin   = Inches(left)
        sec.right_margin  = Inches(right)


# ══════════════════════════════════════════════════════════════════════════════
# CLEAN — minimal monochrome, left-aligned, Calibri, subtle rules
# ══════════════════════════════════════════════════════════════════════════════

def generate_clean():
    doc = Document()
    _set_margins(doc, top=0.75, bottom=0.75, left=0.85, right=0.85)

    # ── Name ──────────────────────────────────────────────────────────────────
    p = _para(doc, "", space_before=0, space_after=2)
    _run(p, "{{NAME}}", "Calibri", 20, bold=True, color=BLACK)

    # ── Contact ───────────────────────────────────────────────────────────────
    p = _para(doc, "", space_before=0, space_after=6)
    _run(p, "{{EMAIL}}", "Calibri", 9, color=MID_GRAY)
    _run(p, "  |  ",    "Calibri", 9, color=LIGHT_GRAY)
    _run(p, "{{PHONE}}","Calibri", 9, color=MID_GRAY)
    _run(p, "  |  ",    "Calibri", 9, color=LIGHT_GRAY)
    _run(p, "{{LOCATION}}", "Calibri", 9, color=MID_GRAY)
    _run(p, "  |  ",    "Calibri", 9, color=LIGHT_GRAY)
    _run(p, "{{LINKEDIN}}", "Calibri", 9, color=MID_GRAY)

    # ── Summary section ───────────────────────────────────────────────────────
    _hr(doc, color=LIGHT_GRAY, thickness_pt=0.5, space_before=0, space_after=4)
    p = _para(doc, "SUMMARY", "Calibri", 9, bold=True, color=BLACK, space_before=0, space_after=2)
    _hr(doc, color=LIGHT_GRAY, thickness_pt=0.5, space_before=0, space_after=3)
    _para(doc, "{{SUMMARY}}", "Calibri", 10, color=DARK_GRAY, space_before=0, space_after=10)

    # ── Experience ────────────────────────────────────────────────────────────
    p = _para(doc, "EXPERIENCE", "Calibri", 9, bold=True, color=BLACK, space_before=0, space_after=2)
    _hr(doc, color=LIGHT_GRAY, thickness_pt=0.5, space_before=0, space_after=3)
    _para(doc, "{{EXPERIENCE}}", "Calibri", 10, color=DARK_GRAY, space_before=0, space_after=10)

    # ── Education ─────────────────────────────────────────────────────────────
    p = _para(doc, "EDUCATION", "Calibri", 9, bold=True, color=BLACK, space_before=0, space_after=2)
    _hr(doc, color=LIGHT_GRAY, thickness_pt=0.5, space_before=0, space_after=3)
    _para(doc, "{{EDUCATION}}", "Calibri", 10, color=DARK_GRAY, space_before=0, space_after=10)

    # ── Additional sections (skills, certs, etc.) ─────────────────────────────
    p = _para(doc, "ADDITIONAL", "Calibri", 9, bold=True, color=BLACK, space_before=0, space_after=2)
    _hr(doc, color=LIGHT_GRAY, thickness_pt=0.5, space_before=0, space_after=3)
    _para(doc, "{{SECTIONS}}", "Calibri", 10, color=DARK_GRAY, space_before=0, space_after=4)

    path = os.path.join(OUT_DIR, "clean.docx")
    doc.save(path)
    print(f"✓  clean.docx  →  {path}")


# ══════════════════════════════════════════════════════════════════════════════
# MODERN — brand-blue accents, Calibri, bold colour headings, contemporary
# ══════════════════════════════════════════════════════════════════════════════

def generate_modern():
    doc = Document()
    _set_margins(doc, top=0.7, bottom=0.7, left=0.85, right=0.85)

    # ── Name — large blue ─────────────────────────────────────────────────────
    p = _para(doc, "", space_before=0, space_after=1)
    _run(p, "{{NAME}}", "Calibri", 26, bold=True, color=BRAND)

    # ── Tagline / contact with teal accent dots ───────────────────────────────
    p = _para(doc, "", space_before=0, space_after=2)
    _run(p, "{{EMAIL}}", "Calibri", 9, color=TEAL)
    _run(p, "  ·  ",     "Calibri", 9, color=LIGHT_GRAY)
    _run(p, "{{PHONE}}", "Calibri", 9, color=TEAL)
    _run(p, "  ·  ",     "Calibri", 9, color=LIGHT_GRAY)
    _run(p, "{{LOCATION}}", "Calibri", 9, color=TEAL)
    _run(p, "  ·  ",     "Calibri", 9, color=LIGHT_GRAY)
    _run(p, "{{LINKEDIN}}", "Calibri", 9, color=TEAL)
    _run(p, "  ·  ",     "Calibri", 9, color=LIGHT_GRAY)
    _run(p, "{{GITHUB}}", "Calibri", 9, color=TEAL)

    # Thick brand rule under header
    _hr(doc, color=BRAND, thickness_pt=2.0, space_before=4, space_after=6)

    def modern_section(title: str, placeholder: str, after: float = 10):
        p = _para(doc, title, "Calibri", 11, bold=True, color=BRAND, space_before=0, space_after=1)
        _hr(doc, color=BRAND, thickness_pt=1.0, space_before=0, space_after=3)
        _para(doc, placeholder, "Calibri", 10, color=DARK_GRAY, space_before=0, space_after=after)

    modern_section("Professional Summary", "{{SUMMARY}}")
    modern_section("Experience", "{{EXPERIENCE}}")
    modern_section("Education", "{{EDUCATION}}")
    modern_section("Skills & Certifications", "{{SECTIONS}}", after=4)

    path = os.path.join(OUT_DIR, "modern.docx")
    doc.save(path)
    print(f"✓  modern.docx  →  {path}")


# ══════════════════════════════════════════════════════════════════════════════
# EXECUTIVE — Georgia serif, centered header, double-rule sections, formal
# ══════════════════════════════════════════════════════════════════════════════

def generate_executive():
    doc = Document()
    _set_margins(doc, top=0.9, bottom=0.9, left=1.0, right=1.0)
    CENTER = WD_ALIGN_PARAGRAPH.CENTER

    # ── Decorative top rule ───────────────────────────────────────────────────
    _hr(doc, color=BLACK, thickness_pt=2.0, space_before=0, space_after=4)

    # ── Name — large, centered, serif ─────────────────────────────────────────
    p = _para(doc, "", align=CENTER, space_before=0, space_after=2)
    _run(p, "{{NAME}}", "Georgia", 24, bold=True, color=BLACK)

    # ── Contact — centered, muted ─────────────────────────────────────────────
    p = _para(doc, "", align=CENTER, space_before=0, space_after=2)
    _run(p, "{{EMAIL}}", "Georgia", 9, color=MID_GRAY)
    _run(p, "  —  ",    "Georgia", 9, color=LIGHT_GRAY)
    _run(p, "{{PHONE}}","Georgia", 9, color=MID_GRAY)
    _run(p, "  —  ",    "Georgia", 9, color=LIGHT_GRAY)
    _run(p, "{{LOCATION}}", "Georgia", 9, color=MID_GRAY)
    _run(p, "  —  ",    "Georgia", 9, color=LIGHT_GRAY)
    _run(p, "{{LINKEDIN}}", "Georgia", 9, color=MID_GRAY)

    # ── Bottom decorative rule ────────────────────────────────────────────────
    _hr(doc, color=BLACK, thickness_pt=2.0, space_before=4, space_after=8)

    def exec_section(title: str, placeholder: str, after: float = 10):
        # Section heading: centered, bold, small caps feel (uppercase Georgia)
        p = _para(doc, title.upper(), "Georgia", 10, bold=True, color=BLACK,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2)
        _double_hr(doc, color=DARK_GRAY, space_before=0, space_after=4)
        _para(doc, placeholder, "Georgia", 10, color=DARK_GRAY, space_before=0, space_after=after)

    exec_section("Professional Summary", "{{SUMMARY}}")
    exec_section("Professional Experience", "{{EXPERIENCE}}")
    exec_section("Education", "{{EDUCATION}}")
    exec_section("Skills & Additional", "{{SECTIONS}}", after=4)

    path = os.path.join(OUT_DIR, "executive.docx")
    doc.save(path)
    print(f"✓  executive.docx  →  {path}")


# ── Main ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    generate_clean()
    generate_modern()
    generate_executive()
    print("\nAll templates generated successfully.")
