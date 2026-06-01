import io
import os
import re
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem,
    Table, TableStyle,
)
from reportlab.platypus import KeepTogether


# ── Shared helpers ─────────────────────────────────────────────────────────────

def _kw_pattern(keywords: list[str]):
    """Compile a case-insensitive regex that matches any keyword (longest first)."""
    if not keywords:
        return None
    terms = sorted(set(k.strip() for k in keywords if k.strip()), key=len, reverse=True)
    return re.compile("|".join(re.escape(t) for t in terms), re.IGNORECASE)


def _bold_pdf(text: str, kw_re) -> str:
    """Wrap keyword matches in <b> tags for ReportLab XML markup.

    Escapes the full text first so XML special chars are safe, then inserts
    <b> tags (which must NOT be escaped) around matched keywords.
    """
    import html as _html_mod
    escaped = _html_mod.escape(str(text or ""))
    if kw_re is None:
        return escaped
    def _wrap(m):
        return f"<b>{_html_mod.escape(m.group(0))}</b>"
    return kw_re.sub(_wrap, escaped)


def _is_url(text: str) -> bool:
    return bool(text and re.match(r'https?://', text.strip()))

def _display_url(url: str) -> str:
    """Strip scheme for display but keep path."""
    return re.sub(r'^https?://', '', url.strip()).rstrip('/')


# ── DOCX helpers ───────────────────────────────────────────────────────────────

def _add_hyperlink(paragraph, url: str, display: str, font_size_pt: int = 10):
    """Insert a clickable, underlined, blue hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run_elem = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')

    # Blue colour
    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), '0563C1')
    rpr.append(color_elem)

    # Single underline
    u_elem = OxmlElement('w:u')
    u_elem.set(qn('w:val'), 'single')
    rpr.append(u_elem)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size_pt * 2))
    rpr.append(sz)

    run_elem.append(rpr)
    t_elem = OxmlElement('w:t')
    t_elem.text = display
    run_elem.append(t_elem)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _add_plain_run(paragraph, text: str, font_size_pt: int = 10, bold: bool = False, muted: bool = False):
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size_pt)
    run.bold = bold
    if muted:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _docx_bullet(doc: Document, text: str, font_size_pt: int = 10, kw_re=None):
    """Add a properly hanging-indented bullet paragraph with optional keyword bolding."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(2)
    # Bullet marker
    bullet_run = p.add_run("•  ")
    bullet_run.font.size = Pt(font_size_pt)
    # Body text — split around keyword matches for bold highlighting
    if kw_re is None:
        r = p.add_run(text)
        r.font.size = Pt(font_size_pt)
    else:
        last = 0
        for m in kw_re.finditer(text):
            if m.start() > last:
                r = p.add_run(text[last:m.start()])
                r.font.size = Pt(font_size_pt)
            br = p.add_run(m.group(0))
            br.bold = True
            br.font.size = Pt(font_size_pt)
            last = m.end()
        if last < len(text):
            r = p.add_run(text[last:])
            r.font.size = Pt(font_size_pt)
    return p


def _docx_section_heading(doc: Document, title: str):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(2)
    run = h.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)  # slate-800 — consistent across all templates

    # Horizontal rule — border via paragraph bottom border (renders as a clean line)
    p = h._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")        # 0.5pt
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "374151")  # slate-700, consistent with heading
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── DOCX generation ────────────────────────────────────────────────────────────

def generate_docx(resume_data: dict, template_path: str, bold_keywords: list[str] | None = None) -> bytes:
    kw_re = _kw_pattern(bold_keywords or [])
    if os.path.exists(template_path):
        return _apply_to_template(resume_data, template_path)
    return _generate_clean_docx(resume_data, kw_re=kw_re)


def _apply_to_template(resume_data: dict, template_path: str) -> bytes:
    doc = Document(template_path)
    replacements = _build_replacements(resume_data)

    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _replace_in_paragraph(para, replacements: dict):
    for key, value in replacements.items():
        if key in para.text:
            for run in para.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)


def _build_replacements(r: dict) -> dict:
    contact = r.get("contact", {})
    exp_lines = []
    for job in r.get("experience", []):
        exp_lines.append(f"{job['role']} at {job['company']} ({job['dates']})")
        for bullet in job.get("bullets", []):
            exp_lines.append(f"  •  {bullet}")
    edu_lines = []
    for ed in r.get("education", []):
        edu_lines.append(f"{ed['degree']} — {ed['institution']} ({ed['dates']})")

    # Build {{SECTIONS}} replacement — works for both new (sections[]) and old
    # (skills[], certifications[]) resume formats so all templates use one token.
    section_lines = []
    if r.get("sections"):
        for sec in r["sections"]:
            title = sec.get("title", "")
            items = sec.get("items", [])
            if not title or not items:
                continue
            block = "\n".join(f"  •  {item}" for item in items)
            section_lines.append(f"{title.upper()}\n{block}")
    else:
        # Backward compat: old-format resumes with flat skills/certifications arrays
        if r.get("skills"):
            block = "\n".join(f"  •  {s}" for s in r["skills"])
            section_lines.append(f"SKILLS\n{block}")
        if r.get("certifications"):
            block = "\n".join(f"  •  {c}" for c in r["certifications"])
            section_lines.append(f"CERTIFICATIONS\n{block}")

    section_map: dict[str, str] = {}
    for sec in r.get("sections", []):
        title = sec.get("title", "")
        items = sec.get("items", [])
        if title and items:
            token = "{{" + title.upper().replace(" ", "_") + "}}"
            section_map[token] = "\n".join(items)

    replacements = {
        "{{NAME}}": r.get("name", ""),
        "{{EMAIL}}": contact.get("email", ""),
        "{{PHONE}}": contact.get("phone", ""),
        "{{LINKEDIN}}": contact.get("linkedin", ""),
        "{{GITHUB}}": contact.get("github", ""),
        "{{WEBSITE}}": contact.get("website", ""),
        "{{LOCATION}}": contact.get("location", ""),
        "{{SUMMARY}}": r.get("summary", ""),
        "{{EXPERIENCE}}": "\n".join(exp_lines),
        "{{EDUCATION}}": "\n".join(edu_lines),
        "{{SECTIONS}}": "\n\n".join(section_lines),
        # Backward compat with old-format sessions
        "{{SKILLS}}": ", ".join(r.get("skills", [])),
        "{{CERTIFICATIONS}}": ", ".join(r.get("certifications", [])),
    }
    replacements.update(section_map)
    return replacements


def _generate_clean_docx(r: dict, kw_re=None) -> bytes:
    doc = Document()
    # A4 page size + narrow margins — ensures page count matches what templates declare
    section = doc.sections[0]
    section.page_width    = Mm(210)
    section.page_height   = Mm(297)
    section.left_margin   = Mm(15)
    section.right_margin  = Mm(15)
    section.top_margin    = Mm(14)
    section.bottom_margin = Mm(12)
    contact = r.get("contact", {})

    # ── Name ──────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(r.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_para.paragraph_format.space_after = Pt(4)

    # ── Contact line — URLs are clickable hyperlinks ───────────────────────────
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(8)

    contact_items: list[tuple[str, str | None]] = []  # (display_text, url_or_None)
    if contact.get("email"):
        contact_items.append((contact["email"], f"mailto:{contact['email']}"))
    if contact.get("phone"):
        contact_items.append((contact["phone"], None))
    if contact.get("location"):
        contact_items.append((contact["location"], None))
    if contact.get("linkedin"):
        url = contact["linkedin"] if _is_url(contact["linkedin"]) else f"https://{contact['linkedin']}"
        contact_items.append((_display_url(url), url))
    if contact.get("github"):
        url = contact["github"] if _is_url(contact["github"]) else f"https://{contact['github']}"
        contact_items.append((_display_url(url), url))
    if contact.get("website"):
        url = contact["website"] if _is_url(contact["website"]) else f"https://{contact['website']}"
        contact_items.append((_display_url(url), url))

    for i, (display, url) in enumerate(contact_items):
        if i > 0:
            _add_plain_run(contact_para, "  |  ", font_size_pt=9, muted=True)
        if url:
            _add_hyperlink(contact_para, url, display, font_size_pt=9)
        else:
            _add_plain_run(contact_para, display, font_size_pt=9, muted=True)

    # ── Summary ────────────────────────────────────────────────────────────────
    if r.get("summary"):
        _docx_section_heading(doc, "Professional Summary")
        p = doc.add_paragraph(r["summary"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10)

    # ── Experience ─────────────────────────────────────────────────────────────
    if r.get("experience"):
        _docx_section_heading(doc, "Experience")
        for job in r["experience"]:
            role_para = doc.add_paragraph()
            role_para.paragraph_format.space_after = Pt(1)
            role_run = role_para.add_run(f"{job['role']}")
            role_run.bold = True
            role_run.font.size = Pt(10.5)
            role_para.add_run(f"  —  {job['company']}")
            role_para.runs[-1].font.size = Pt(10.5)

            dates_para = doc.add_paragraph(job.get("dates", ""))
            dates_para.paragraph_format.space_after = Pt(3)
            for run in dates_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                run.italic = True

            for bullet in job.get("bullets", []):
                _docx_bullet(doc, bullet, kw_re=kw_re)

            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Education ──────────────────────────────────────────────────────────────
    if r.get("education"):
        _docx_section_heading(doc, "Education")
        for ed in r["education"]:
            ed_para = doc.add_paragraph()
            ed_para.paragraph_format.space_after = Pt(1)
            deg_run = ed_para.add_run(f"{ed['degree']}")
            deg_run.bold = True
            deg_run.font.size = Pt(10.5)
            ed_para.add_run(f"  —  {ed['institution']}")
            ed_para.runs[-1].font.size = Pt(10.5)

            dates_para = doc.add_paragraph(ed.get("dates", ""))
            dates_para.paragraph_format.space_after = Pt(4)
            for run in dates_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                run.italic = True

    # ── Additional sections — dynamic (new format) or legacy fallback ──────────
    if r.get("sections"):
        for sec in r["sections"]:
            title = sec.get("title", "")
            items = sec.get("items", [])
            if not title or not items:
                continue
            _docx_section_heading(doc, title)
            for item in items:
                _docx_bullet(doc, item, kw_re=kw_re)
    else:
        # Backward compat: old-format resumes stored before dynamic sections
        if r.get("skills"):
            _docx_section_heading(doc, "Skills")
            for skill in r["skills"]:
                _docx_bullet(doc, skill, kw_re=kw_re)
        if r.get("certifications"):
            _docx_section_heading(doc, "Certifications")
            for cert in r["certifications"]:
                _docx_bullet(doc, cert, kw_re=kw_re)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF generation (pure Python, no LibreOffice) ─────────────────────────────

import html as _html

_BRAND  = colors.HexColor("#2B579A")
_TEXT   = colors.HexColor("#1a1a1a")
_MUTED  = colors.HexColor("#555555")
_LINK   = colors.HexColor("#0563C1")


def _e(text: str) -> str:
    """Escape HTML entities so ReportLab's XML parser doesn't choke."""
    return _html.escape(str(text or ""))


def _pdf_link(url: str, display: str) -> str:
    """Return an inline ReportLab hyperlink tag."""
    return f'<a href="{_e(url)}" color="#0563C1"><u>{_e(display)}</u></a>'


def generate_pdf(resume_data: dict, bold_keywords: list[str] | None = None) -> bytes:
    kw_re = _kw_pattern(bold_keywords or [])
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=14*mm, rightMargin=14*mm,
        topMargin=12*mm, bottomMargin=12*mm,
    )

    base = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "Name", parent=base["Normal"],
        fontSize=22, leading=26, textColor=_TEXT,
        fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "Contact", parent=base["Normal"],
        fontSize=9, leading=12, textColor=_MUTED,
        fontName="Helvetica", alignment=TA_CENTER, spaceAfter=6,
    )
    section_style = ParagraphStyle(
        "Section", parent=base["Normal"],
        fontSize=10, leading=12, textColor=_BRAND,
        fontName="Helvetica-Bold", spaceBefore=8, spaceAfter=2,
    )
    job_title_style = ParagraphStyle(
        "JobTitle", parent=base["Normal"],
        fontSize=10, leading=12, textColor=_TEXT,
        fontName="Helvetica-Bold", spaceAfter=1,
    )
    job_meta_style = ParagraphStyle(
        "JobMeta", parent=base["Normal"],
        fontSize=9, leading=11, textColor=_MUTED,
        fontName="Helvetica-Oblique", spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "Body", parent=base["Normal"],
        fontSize=9.5, leading=13, textColor=_TEXT,
        fontName="Helvetica", spaceAfter=4,
        alignment=TA_JUSTIFY,
    )
    # Bullet text style — no indent here; alignment is handled by the Table column
    bullet_text_style = ParagraphStyle(
        "BulletText", parent=base["Normal"],
        fontSize=9.5, leading=13, textColor=_TEXT,
        fontName="Helvetica", spaceAfter=0,
    )

    def hr():
        return HRFlowable(width="100%", thickness=0.5, color=_BRAND, spaceAfter=3)

    def section_heading(title: str):
        return [Paragraph(title.upper(), section_style), hr()]

    def _pdf_bullet_table(bullets: list[str]) -> Table:
        """Render bullets as a 2-column table: marker | text.
        This guarantees wrapped lines align with the start of text, not the bullet marker.
        """
        marker_style = ParagraphStyle(
            "BulletMarker", parent=base["Normal"],
            fontSize=9.5, leading=13, textColor=_TEXT,
            fontName="Helvetica",
        )
        rows = [
            [
                Paragraph("•", marker_style),
                Paragraph(_bold_pdf(b, kw_re), bullet_text_style),
            ]
            for b in bullets
        ]
        tbl = Table(rows, colWidths=[10, None])
        tbl.setStyle(TableStyle([
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING",   (0, 0), (-1, -1), 0),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
            ("TOPPADDING",    (0, 0), (-1, -1), 1),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
        ]))
        return tbl

    story = []
    r = resume_data
    contact = r.get("contact", {})

    # ── Name ──────────────────────────────────────────────────────────────────
    story.append(Paragraph(_e(r.get("name", "")), name_style))

    # ── Contact — clickable URLs ───────────────────────────────────────────────
    contact_parts = []
    if contact.get("email"):
        contact_parts.append(_pdf_link(f"mailto:{contact['email']}", contact["email"]))
    if contact.get("phone"):
        contact_parts.append(_e(contact["phone"]))
    if contact.get("location"):
        contact_parts.append(_e(contact["location"]))
    if contact.get("linkedin"):
        url = contact["linkedin"] if _is_url(contact["linkedin"]) else f"https://{contact['linkedin']}"
        contact_parts.append(_pdf_link(url, _display_url(url)))
    if contact.get("github"):
        url = contact["github"] if _is_url(contact["github"]) else f"https://{contact['github']}"
        contact_parts.append(_pdf_link(url, _display_url(url)))
    if contact.get("website"):
        url = contact["website"] if _is_url(contact["website"]) else f"https://{contact['website']}"
        contact_parts.append(_pdf_link(url, _display_url(url)))

    if contact_parts:
        story.append(Paragraph("  ·  ".join(contact_parts), contact_style))

    story.append(HRFlowable(width="100%", thickness=1, color=_BRAND, spaceAfter=5))

    # ── Summary ────────────────────────────────────────────────────────────────
    if r.get("summary"):
        story += section_heading("Professional Summary")
        story.append(Paragraph(_bold_pdf(r["summary"], kw_re), body_style))

    # ── Experience ─────────────────────────────────────────────────────────────
    if r.get("experience"):
        story += section_heading("Experience")
        for job in r["experience"]:
            block = [
                Paragraph(f"<b>{_e(job['role'])}</b>  ·  {_e(job['company'])}", job_title_style),
                Paragraph(_e(job.get("dates", "")), job_meta_style),
            ]
            bullets = job.get("bullets", [])
            if bullets:
                block.append(_pdf_bullet_table(bullets))
            block.append(Spacer(1, 3))
            story.append(KeepTogether(block))

    # ── Education ──────────────────────────────────────────────────────────────
    if r.get("education"):
        story += section_heading("Education")
        for ed in r["education"]:
            story.append(Paragraph(
                f"<b>{_e(ed['degree'])}</b>  ·  {_e(ed['institution'])}",
                job_title_style,
            ))
            story.append(Paragraph(_e(ed.get("dates", "")), job_meta_style))

    # ── Additional sections — dynamic (new format) or legacy fallback ──────────
    if r.get("sections"):
        for sec in r["sections"]:
            title = sec.get("title", "")
            items = sec.get("items", [])
            if not title or not items:
                continue
            story += section_heading(title)
            story.append(_pdf_bullet_table(items))
    else:
        # Backward compat: old-format resumes stored before dynamic sections
        if r.get("skills"):
            story += section_heading("Skills")
            story.append(_pdf_bullet_table(r["skills"]))
        if r.get("certifications"):
            story += section_heading("Certifications")
            story.append(_pdf_bullet_table(r["certifications"]))

    doc.build(story)
    return buf.getvalue()
