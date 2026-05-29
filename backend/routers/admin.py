"""Admin-only router.

All endpoints require is_superadmin=True on the authenticated user.

GET  /api/admin/users              — all users with activity stats
GET  /api/admin/audit              — paginated audit log
GET  /api/admin/prompts            — current prompt values (override or default)
PUT  /api/admin/prompts/{key}      — set a prompt override
DELETE /api/admin/prompts/{key}    — remove override (revert to hardcoded default)
"""
from __future__ import annotations

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from fastapi.responses import FileResponse
from pydantic import BaseModel
from bson import ObjectId

from dependencies.auth import require_superadmin
from database import get_db
from services.prompt_store import PROMPT_KEYS, get_override, set_override, delete_override, list_overrides

# Import hardcoded defaults so admin UI can show them when no override exists
from services.pipeline.prompts.anthropic import (
    _GENERATOR_SYSTEM_BASE,
    _JOB_ANALYZER_SYSTEM,
    _ANTHROPIC_EVALUATOR_BASE,
)
from services.pipeline.prompts.openai import _OPENAI_EVALUATOR_BASE
from services.pipeline.prompts.google import _GOOGLE_EVALUATOR_BASE

DEFAULTS: dict[str, str] = {
    "generator_system": _GENERATOR_SYSTEM_BASE,
    "job_analyzer_system": _JOB_ANALYZER_SYSTEM,
    "anthropic_evaluator_base": _ANTHROPIC_EVALUATOR_BASE,
    "openai_evaluator_base": _OPENAI_EVALUATOR_BASE,
    "google_evaluator_base": _GOOGLE_EVALUATOR_BASE,
}

router = APIRouter()


# ── Users ─────────────────────────────────────────────────────────────────────

@router.get("/admin/users")
async def list_users(_: dict = Depends(require_superadmin)):
    """Return all users — basic info only. Fetch per-user stats separately via /admin/users/{id}/stats."""
    db = get_db()
    users = await db.users.find(
        {},
        {"hashed_password": 0}
    ).sort("created_at", -1).to_list(length=1000)
    return [
        {
            "id": str(u["_id"]),
            "email": u.get("email", ""),
            "name": u.get("name", ""),
            "tier": u.get("tier", "free"),
            "is_active": u.get("is_active", True),
            "is_superadmin": u.get("is_superadmin", False),
            "created_at": u.get("created_at"),
        }
        for u in users
    ]


@router.get("/admin/users/{user_id}/stats")
async def get_user_stats(user_id: str, _: dict = Depends(require_superadmin)):
    """Fetch activity counts for a single user — called lazily when a user row is expanded."""
    import asyncio
    from bson import ObjectId
    db = get_db()
    try:
        oid = ObjectId(user_id)
    except Exception:
        raise HTTPException(400, "Invalid user ID.")
    session_count, resume_count, alert_count, saved_job_count = await asyncio.gather(
        db.sessions.count_documents({"user_id": str(oid)}),
        db.saved_resumes.count_documents({"user_id": oid}),
        db.job_alerts.count_documents({"user_id": oid}),
        db.saved_jobs.count_documents({"user_id": str(oid)}),
    )
    return {
        "user_id": user_id,
        "session_count": session_count,
        "resume_count": resume_count,
        "alert_count": alert_count,
        "saved_job_count": saved_job_count,
    }


# ── Audit log ─────────────────────────────────────────────────────────────────

@router.get("/admin/audit")
async def list_audit(
    page: int = Query(1, ge=1),
    page_size: int = Query(50, ge=1, le=200),
    _: dict = Depends(require_superadmin),
):
    db = get_db()
    skip = (page - 1) * page_size
    total = await db.audit_log.count_documents({})
    docs = await db.audit_log.find({}).sort("created_at", -1).skip(skip).limit(page_size).to_list(length=page_size)
    return {
        "total": total,
        "page": page,
        "page_size": page_size,
        "items": [
            {
                "id": str(d["_id"]),
                "user_id": d.get("user_id", ""),
                "user_email": d.get("user_email", ""),
                "action": d.get("action", ""),
                "metadata": d.get("metadata", {}),
                "created_at": d.get("created_at"),
            }
            for d in docs
        ],
    }


# ── Prompts ───────────────────────────────────────────────────────────────────

@router.get("/admin/prompts")
async def list_prompts(_: dict = Depends(require_superadmin)):
    overrides = await list_overrides()
    return [
        {
            "key": key,
            "label": label,
            "body": overrides.get(key, DEFAULTS.get(key, "")),
            "is_override": key in overrides,
            "default_body": DEFAULTS.get(key, ""),
        }
        for key, label in PROMPT_KEYS.items()
    ]


class PromptBody(BaseModel):
    body: str


@router.put("/admin/prompts/{key}")
async def update_prompt(key: str, payload: PromptBody, _: dict = Depends(require_superadmin)):
    if key not in PROMPT_KEYS:
        raise HTTPException(400, f"Unknown prompt key: {key}")
    if not payload.body.strip():
        raise HTTPException(400, "Prompt body cannot be empty.")
    await set_override(key, payload.body.strip())
    return {"key": key, "saved": True}


@router.delete("/admin/prompts/{key}")
async def reset_prompt(key: str, _: dict = Depends(require_superadmin)):
    if key not in PROMPT_KEYS:
        raise HTTPException(400, f"Unknown prompt key: {key}")
    await delete_override(key)
    return {"key": key, "reset": True, "default_body": DEFAULTS.get(key, "")}


# ── Professions ───────────────────────────────────────────────────────────────

class ProfessionUpsertBody(BaseModel):
    slug: str
    display_name: str
    keywords: list[str] = []
    generator_context: str = ""
    evaluator_context: str = ""
    scoring_criteria: str = ""
    aggregator_context: str = ""
    evaluator_names: list[str] = []


class ProfessionPatchBody(BaseModel):
    display_name: str | None = None
    keywords: list[str] | None = None
    generator_context: str | None = None
    evaluator_context: str | None = None
    scoring_criteria: str | None = None
    aggregator_context: str | None = None
    evaluator_names: list[str] | None = None
    is_active: bool | None = None


@router.get("/admin/professions")
async def admin_list_professions(_: dict = Depends(require_superadmin)):
    """Return all professions (active and inactive) for admin management."""
    db = get_db()
    docs = await db.professions.find({}, {"_id": 0}).sort("slug", 1).to_list(None)
    return docs


@router.post("/admin/professions", status_code=201)
async def admin_create_profession(body: ProfessionUpsertBody, _: dict = Depends(require_superadmin)):
    db = get_db()
    if await db.professions.find_one({"slug": body.slug}):
        raise HTTPException(409, f"Profession '{body.slug}' already exists.")
    from datetime import datetime
    now = datetime.utcnow().isoformat()
    doc = {**body.model_dump(), "is_active": True, "created_at": now, "updated_at": now}
    await db.professions.insert_one(doc)
    doc.pop("_id", None)
    return doc


@router.patch("/admin/professions/{slug}")
async def admin_update_profession(slug: str, body: ProfessionPatchBody, _: dict = Depends(require_superadmin)):
    db = get_db()
    from datetime import datetime
    updates = {k: v for k, v in body.model_dump().items() if v is not None}
    updates["updated_at"] = datetime.utcnow().isoformat()
    result = await db.professions.find_one_and_update(
        {"slug": slug},
        {"$set": updates},
        return_document=True,
    )
    if not result:
        raise HTTPException(404, f"Profession '{slug}' not found.")
    result.pop("_id", None)
    return result


@router.delete("/admin/professions/{slug}", status_code=204)
async def admin_delete_profession(slug: str, _: dict = Depends(require_superadmin)):
    if slug == "generic":
        raise HTTPException(400, "Cannot delete the generic fallback profession.")
    db = get_db()
    from datetime import datetime
    result = await db.professions.update_one(
        {"slug": slug},
        {"$set": {"is_active": False, "updated_at": datetime.utcnow().isoformat()}},
    )
    if result.modified_count == 0:
        raise HTTPException(404, f"Profession '{slug}' not found.")


# ── Templates ─────────────────────────────────────────────────────────────────

def _serialize_template(doc: dict) -> dict:
    return {
        "id": str(doc["_id"]),
        "name": doc.get("name", ""),
        "type": doc.get("type", "custom"),
        "description": doc.get("description", ""),
        "placeholders": doc.get("placeholders", []),
        "preview_image_url": doc.get("preview_image_url", ""),
        "file_path": doc.get("file_path", ""),
        "is_active": doc.get("is_active", True),
        "created_at": doc.get("created_at"),
        "updated_at": doc.get("updated_at"),
    }


@router.get("/admin/templates")
async def admin_list_templates(_: dict = Depends(require_superadmin)):
    """Return ALL templates including inactive ones."""
    db = get_db()
    docs = await db.templates.find({}).sort("created_at", 1).to_list(length=200)
    return [_serialize_template(d) for d in docs]


@router.post("/admin/templates/upload", status_code=201)
async def admin_upload_template(
    file: UploadFile = File(...),
    name: str = Form(...),
    description: str = Form(""),
    _: dict = Depends(require_superadmin),
):
    """Upload a new template DOCX, validate its placeholders, and store it."""
    import os, io
    from datetime import datetime
    from docx import Document

    if not file.filename or not file.filename.endswith(".docx"):
        raise HTTPException(400, "Only .docx files are accepted.")

    file_bytes = await file.read()
    if len(file_bytes) > 5 * 1024 * 1024:
        raise HTTPException(400, "File exceeds 5 MB limit.")

    # Extract placeholders from the DOCX
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + "\n".join(p.text for p in cell.paragraphs)
    except Exception as exc:
        raise HTTPException(422, f"Could not read DOCX: {exc}")

    import re
    found_placeholders = sorted(set(re.findall(r"\{\{[A-Z_]+\}\}", full_text)))

    # Require at minimum the core contact + content placeholders
    required = {"{{NAME}}", "{{SUMMARY}}", "{{EXPERIENCE}}", "{{EDUCATION}}"}
    missing = required - set(found_placeholders)
    if missing:
        raise HTTPException(
            422,
            f"Template is missing required placeholders: {', '.join(sorted(missing))}. "
            f"Found: {', '.join(found_placeholders) or 'none'}",
        )

    # Save file
    uploads_dir = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "templates", "uploads",
    )
    os.makedirs(uploads_dir, exist_ok=True)
    safe_name = f"{datetime.utcnow().timestamp()}_{file.filename}"
    dest = os.path.join(uploads_dir, safe_name)
    with open(dest, "wb") as f_out:
        f_out.write(file_bytes)

    db = get_db()
    now = datetime.utcnow()
    result = await db.templates.insert_one({
        "name": name.strip(),
        "type": "custom",
        "description": description.strip(),
        "preview_image_url": "",
        "file_path": f"templates/uploads/{safe_name}",
        "placeholders": found_placeholders,
        "is_active": True,
        "created_at": now,
        "updated_at": now,
    })
    doc_out = await db.templates.find_one({"_id": result.inserted_id})
    return _serialize_template(doc_out)


class TemplatePatchBody(BaseModel):
    name: str | None = None
    description: str | None = None
    is_active: bool | None = None


@router.patch("/admin/templates/{template_id}")
async def admin_update_template(
    template_id: str,
    body: TemplatePatchBody,
    _: dict = Depends(require_superadmin),
):
    from datetime import datetime
    db = get_db()
    try:
        oid = ObjectId(template_id)
    except Exception:
        raise HTTPException(400, "Invalid template ID.")
    updates = {k: v for k, v in body.model_dump().items() if v is not None}
    if not updates:
        raise HTTPException(400, "No fields to update.")
    updates["updated_at"] = datetime.utcnow()
    result = await db.templates.find_one_and_update(
        {"_id": oid}, {"$set": updates}, return_document=True
    )
    if not result:
        raise HTTPException(404, "Template not found.")
    return _serialize_template(result)


@router.delete("/admin/templates/{template_id}", status_code=204)
async def admin_delete_template(template_id: str, _: dict = Depends(require_superadmin)):
    import os
    from datetime import datetime
    from services.template_service import get_template_path
    db = get_db()
    try:
        oid = ObjectId(template_id)
    except Exception:
        raise HTTPException(400, "Invalid template ID.")
    doc = await db.templates.find_one({"_id": oid})
    if not doc:
        raise HTTPException(404, "Template not found.")
    if doc.get("type") == "prebuilt":
        raise HTTPException(400, "Prebuilt templates cannot be deleted. Deactivate them instead.")
    # Delete the file if it's in uploads
    file_path = doc.get("file_path", "")
    if "uploads" in file_path:
        abs_path = get_template_path(file_path)
        if os.path.exists(abs_path):
            os.remove(abs_path)
    await db.templates.delete_one({"_id": oid})


@router.get("/admin/templates/{template_id}/download")
async def admin_download_template(template_id: str, _: dict = Depends(require_superadmin)):
    from services.template_service import get_template_path
    db = get_db()
    try:
        oid = ObjectId(template_id)
    except Exception:
        raise HTTPException(400, "Invalid template ID.")
    doc = await db.templates.find_one({"_id": oid})
    if not doc:
        raise HTTPException(404, "Template not found.")
    abs_path = get_template_path(doc.get("file_path", ""))
    if not abs_path or not __import__("os").path.exists(abs_path):
        raise HTTPException(404, "Template file not found on disk.")
    filename = f"{doc.get('name', 'template')}.docx"
    return FileResponse(abs_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)
